"""PRD export endpoints: Markdown, DOCX, PDF."""

from __future__ import annotations

import io
import logging
from datetime import datetime, date
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import Response, StreamingResponse

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from xhtml2pdf import pisa

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/export", tags=["export"])

# ──────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────

def _sanitize_filename(name: str) -> str:
    """Remove characters unsafe for filenames."""
    safe = name.replace(" ", "_")
    safe = "".join(c for c in safe if c.isalnum() or c in "._-")
    return safe or "export"

def _format_date(dt: Any) -> str:
    """Format datetime or string to readable date."""
    if isinstance(dt, datetime):
        return dt.strftime("%Y-%m-%d %H:%M")
    if isinstance(dt, date):
        return dt.strftime("%Y-%m-%d")
    if isinstance(dt, str):
        return dt[:19].replace("T", " ")
    return str(dt)

PRIORITY_LABEL: Dict[str, str] = {
    "P0": "P0 — 紧急",
    "P1": "P1 — 高",
    "P2": "P2 — 中",
    "P3": "P3 — 低",
}

# ──────────────────────────────────────────────
# Markdown Export
# ──────────────────────────────────────────────

def export_to_markdown(prd: Dict[str, Any]) -> str:
    """Render a PRD dict to a Markdown string."""
    app_name = prd.get("app_name") or "Unknown App"
    title = prd.get("title") or f"PRD - {app_name}"
    goal = prd.get("analysis_goal") or "N/A"
    generated = _format_date(prd.get("generated_at", datetime.now()))
    background = prd.get("background") or ""
    problem = prd.get("problem_statement") or ""
    findings = prd.get("supporting_findings") or []
    user_stories: List[Dict] = prd.get("user_stories") or []
    requirements: List[Dict] = prd.get("requirements") or []
    version_plan: List[Dict] = prd.get("version_plan") or []

    lines: List[str] = []

    # Header
    lines.append(f"# {title}")
    lines.append("")
    lines.append(f"- **App**: {app_name}")
    lines.append(f"- **Analysis Goal**: {goal}")
    lines.append(f"- **Generated**: {generated}")
    lines.append("")

    # Background
    if background:
        lines.append("## Background")
        lines.append("")
        lines.append(background)
        lines.append("")

    # Problem Statement
    if problem:
        lines.append("## Problem Statement")
        lines.append("")
        lines.append(problem)
        lines.append("")

    # Supporting Findings
    if findings:
        lines.append("## Supporting Findings")
        lines.append("")
        for f in findings:
            lines.append(f"- {f}")
        lines.append("")

    # User Stories
    if user_stories:
        lines.append("## User Stories")
        lines.append("")
        lines.append("| ID | Role | Goal | Benefit |")
        lines.append("|----|------|------|---------|")
        for us in user_stories:
            rid = us.get("id", "")
            role = us.get("role") or us.get("role_en") or ""
            gl = us.get("goal") or us.get("goal_en") or ""
            benefit = us.get("benefit") or us.get("benefit_en") or ""
            lines.append(f"| {rid} | {role} | {gl} | {benefit} |")
        lines.append("")

    # Requirements
    if requirements:
        lines.append("## Requirements")
        lines.append("")
        for i, req in enumerate(requirements, 1):
            rid = req.get("req_id", f"REQ-{i}")
            rtitle = req.get("title", "")
            priority = req.get("priority", "")
            description = req.get("description", "")
            user_problem = req.get("user_problem", "")
            business_value = req.get("business_value", "")
            target = req.get("target_version") or "TBD"
            criteria: List[str] = req.get("acceptance_criteria") or []
            is_assumption = req.get("is_assumption", False)

            lines.append(f"### {rid}: {rtitle}")
            lines.append("")
            lines.append(f"- **Priority**: {PRIORITY_LABEL.get(priority, priority)}")
            lines.append(f"- **Target Version**: {target}")
            if is_assumption:
                lines.append("- **⚠ Assumption**")
            lines.append("")
            lines.append(f"**Description**: {description}")
            lines.append("")
            if user_problem:
                lines.append(f"**User Problem**: {user_problem}")
                lines.append("")
            if business_value:
                lines.append(f"**Business Value**: {business_value}")
                lines.append("")
            if criteria:
                lines.append("**Acceptance Criteria**:")
                for c in criteria:
                    lines.append(f"- {c}")
                lines.append("")
            lines.append("---")
            lines.append("")

    # Version Plan
    if version_plan:
        lines.append("## Version Plan")
        lines.append("")
        for vp in version_plan:
            ver = vp.get("version", "Unknown")
            theme = vp.get("theme", "")
            release_goal = vp.get("release_goal", "")
            req_ids: List[str] = vp.get("requirement_ids") or []
            rationale = vp.get("rationale", "")

            lines.append(f"### {ver} — {theme}")
            lines.append("")
            lines.append(f"- **Goal**: {release_goal}")
            lines.append(f"- **Requirements**: {', '.join(req_ids) if req_ids else 'None'}")
            lines.append(f"- **Rationale**: {rationale}")
            lines.append("")

    # Footer
    lines.append("---")
    lines.append(f"*Generated by App Review Insights on {generated}*")
    lines.append("")

    return "\n".join(lines)


# ──────────────────────────────────────────────
# DOCX Export
# ──────────────────────────────────────────────

def _add_docx_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # blue-800

def _add_docx_meta(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    p.add_run(value).font.size = Pt(10)

def _add_docx_table(doc: Document, headers: List[str], rows: List[List[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r, row in enumerate(rows, 1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


def export_to_docx(prd: Dict[str, Any]) -> bytes:
    """Render a PRD dict to a .docx file (bytes)."""
    app_name = prd.get("app_name") or "Unknown App"
    title = prd.get("title") or f"PRD - {app_name}"
    goal = prd.get("analysis_goal") or "N/A"
    generated = _format_date(prd.get("generated_at", datetime.now()))
    background = prd.get("background") or ""
    problem = prd.get("problem_statement") or ""
    findings = prd.get("supporting_findings") or []
    user_stories: List[Dict] = prd.get("user_stories") or []
    requirements: List[Dict] = prd.get("requirements") or []
    version_plan: List[Dict] = prd.get("version_plan") or []

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    # Divider
    doc.add_paragraph("─" * 60)

    # Meta
    _add_docx_meta(doc, "App", app_name)
    _add_docx_meta(doc, "Analysis Goal", goal)
    _add_docx_meta(doc, "Generated", generated)
    doc.add_paragraph()

    # Background
    if background:
        _add_docx_heading(doc, "Background", level=2)
        doc.add_paragraph(background)

    # Problem Statement
    if problem:
        _add_docx_heading(doc, "Problem Statement", level=2)
        doc.add_paragraph(problem)

    # Supporting Findings
    if findings:
        _add_docx_heading(doc, "Supporting Findings", level=2)
        for f in findings:
            doc.add_paragraph(f, style="List Bullet")

    # User Stories
    if user_stories:
        _add_docx_heading(doc, "User Stories", level=2)
        _add_docx_table(
            doc,
            ["ID", "Role", "Goal", "Benefit"],
            [
                [
                    us.get("id", ""),
                    us.get("role") or us.get("role_en") or "",
                    us.get("goal") or us.get("goal_en") or "",
                    us.get("benefit") or us.get("benefit_en") or "",
                ]
                for us in user_stories
            ],
        )
        doc.add_paragraph()

    # Requirements
    if requirements:
        _add_docx_heading(doc, "Requirements", level=2)
        for i, req in enumerate(requirements, 1):
            rid = req.get("req_id", f"REQ-{i}")
            rtitle = req.get("title", "")
            priority = req.get("priority", "")
            description = req.get("description", "")
            user_problem = req.get("user_problem", "")
            business_value = req.get("business_value", "")
            target = req.get("target_version") or "TBD"
            criteria: List[str] = req.get("acceptance_criteria") or []
            is_assumption = req.get("is_assumption", False)

            _add_docx_heading(doc, f"{rid}: {rtitle}", level=3)
            _add_docx_meta(doc, "Priority", PRIORITY_LABEL.get(priority, priority))
            _add_docx_meta(doc, "Target Version", target)
            if is_assumption:
                p = doc.add_paragraph()
                run = p.add_run("⚠ Assumption - needs validation")
                run.italic = True
                run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

            doc.add_paragraph(f"Description: {description}")
            if user_problem:
                doc.add_paragraph(f"User Problem: {user_problem}")
            if business_value:
                doc.add_paragraph(f"Business Value: {business_value}")
            if criteria:
                p = doc.add_paragraph()
                p.add_run("Acceptance Criteria:").bold = True
                for c in criteria:
                    doc.add_paragraph(c, style="List Bullet")

    # Version Plan
    if version_plan:
        _add_docx_heading(doc, "Version Plan", level=2)
        for vp in version_plan:
            ver = vp.get("version", "Unknown")
            theme = vp.get("theme", "")
            release_goal = vp.get("release_goal", "")
            req_ids: List[str] = vp.get("requirement_ids") or []
            rationale = vp.get("rationale", "")

            _add_docx_heading(doc, f"{ver} — {theme}", level=3)
            _add_docx_meta(doc, "Goal", release_goal)
            _add_docx_meta(doc, "Requirements", ", ".join(req_ids) if req_ids else "None")
            doc.add_paragraph(f"Rationale: {rationale}")

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(f"Generated by App Review Insights on {generated}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)  # slate-400

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ──────────────────────────────────────────────
# PDF Export (via reportlab)
# ──────────────────────────────────────────────

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.colors import HexColor
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    ListFlowable, ListItem, HRFlowable, KeepTogether,
)

# ── Styles ──

_BASE_STYLES = getSampleStyleSheet()

STYLE_H1 = ParagraphStyle("H1_Custom", parent=_BASE_STYLES["h1"],
    fontSize=20, textColor=HexColor("#1e3a8a"), spaceAfter=6, alignment=TA_CENTER)
STYLE_H2 = ParagraphStyle("H2_Custom", parent=_BASE_STYLES["h2"],
    fontSize=14, textColor=HexColor("#1e40af"), spaceBefore=18, spaceAfter=6,
    borderPadding=(0, 0, 2, 0))
STYLE_H3 = ParagraphStyle("H3_Custom", parent=_BASE_STYLES["h3"],
    fontSize=11, textColor=HexColor("#334155"), spaceBefore=12, spaceAfter=4)
STYLE_BODY = ParagraphStyle("Body_Custom", parent=_BASE_STYLES["Normal"],
    fontSize=10, textColor=HexColor("#1e293b"), leading=16, spaceAfter=6)
STYLE_META = ParagraphStyle("Meta", parent=_BASE_STYLES["Normal"],
    fontSize=9, textColor=HexColor("#64748b"), spaceAfter=2)
STYLE_FOOTER = ParagraphStyle("Footer", parent=_BASE_STYLES["Normal"],
    fontSize=8, textColor=HexColor("#94a3b8"), alignment=TA_CENTER)
STYLE_TABLE_HEADER = ParagraphStyle("TH", parent=_BASE_STYLES["Normal"],
    fontSize=8, textColor=HexColor("#ffffff"), fontName="Helvetica-Bold")
STYLE_TABLE_CELL = ParagraphStyle("TD", parent=_BASE_STYLES["Normal"],
    fontSize=8, textColor=HexColor("#1e293b"))
STYLE_BULLET = ParagraphStyle("Bullet", parent=_BASE_STYLES["Normal"],
    fontSize=9, textColor=HexColor("#475569"), leftIndent=16, bulletIndent=8,
    spaceBefore=1, spaceAfter=1)


def _build_pdf_story(prd: Dict[str, Any]) -> list:
    """Build a reportlab story (list of flowables) from PRD dict."""
    app_name = prd.get("app_name") or "Unknown App"
    title = prd.get("title") or f"PRD - {app_name}"
    goal = prd.get("analysis_goal") or "N/A"
    generated = _format_date(prd.get("generated_at", datetime.now()))
    background = prd.get("background") or ""
    problem = prd.get("problem_statement") or ""
    findings = prd.get("supporting_findings") or []
    user_stories: List[Dict] = prd.get("user_stories") or []
    requirements: List[Dict] = prd.get("requirements") or []
    version_plan: List[Dict] = prd.get("version_plan") or []

    story: list = []

    # Title
    story.append(Paragraph(title, STYLE_H1))
    story.append(Spacer(1, 4))

    # Meta
    story.append(Paragraph(f"<b>App:</b> {app_name}", STYLE_META))
    story.append(Paragraph(f"<b>Analysis Goal:</b> {goal}", STYLE_META))
    story.append(Paragraph(f"<b>Generated:</b> {generated}", STYLE_META))
    story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#cbd5e1"),
        spaceBefore=8, spaceAfter=12))

    # Background
    if background:
        story.append(Paragraph("Background", STYLE_H2))
        story.append(Paragraph(background, STYLE_BODY))

    # Problem Statement
    if problem:
        story.append(Paragraph("Problem Statement", STYLE_H2))
        story.append(Paragraph(problem, STYLE_BODY))

    # Supporting Findings
    if findings:
        story.append(Paragraph("Supporting Findings", STYLE_H2))
        items = [ListItem(Paragraph(f, STYLE_BULLET)) for f in findings]
        story.append(ListFlowable(items, bulletType="bullet", start="- "))

    # User Stories
    if user_stories:
        story.append(Paragraph("User Stories", STYLE_H2))
        header = [Paragraph(h, STYLE_TABLE_HEADER) for h in ["ID", "Role", "Goal", "Benefit"]]
        data = [header]
        for us in user_stories:
            data.append([
                Paragraph(us.get("id", ""), STYLE_TABLE_CELL),
                Paragraph(us.get("role") or us.get("role_en") or "", STYLE_TABLE_CELL),
                Paragraph(us.get("goal") or us.get("goal_en") or "", STYLE_TABLE_CELL),
                Paragraph(us.get("benefit") or us.get("benefit_en") or "", STYLE_TABLE_CELL),
            ])
        col_w = [A4[0]*0.1, A4[0]*0.25, A4[0]*0.3, A4[0]*0.25]
        tbl = Table(data, colWidths=col_w, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), HexColor("#1e40af")),
            ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e2e8f0")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#ffffff"), HexColor("#f8fafc")]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(tbl)

    # Requirements
    if requirements:
        story.append(Paragraph("Requirements", STYLE_H2))
        for i, req in enumerate(requirements):
            rid = req.get("req_id", f"REQ-{i+1}")
            rtitle = req.get("title", "")
            priority = req.get("priority", "")
            target = req.get("target_version") or "TBD"
            desc = req.get("description", "")
            user_problem = req.get("user_problem", "")
            business_value = req.get("business_value", "")
            criteria: List = req.get("acceptance_criteria") or []
            is_assumption = req.get("is_assumption", False)

            meta_parts = [f"<b>Priority:</b> {PRIORITY_LABEL.get(priority, priority)}",
                          f"<b>Target:</b> {target}"]
            if is_assumption:
                meta_parts.append('<font color="#b45309"><i>⚠ Assumption</i></font>')
            meta_line = " &nbsp;|&nbsp; ".join(meta_parts)

            elems: list = [
                Paragraph(f"{rid}: {rtitle}", STYLE_H3),
                Paragraph(meta_line, STYLE_META),
                Paragraph(f"<b>Description:</b> {desc}", STYLE_BODY),
            ]
            if user_problem:
                elems.append(Paragraph(f"<b>User Problem:</b> {user_problem}", STYLE_BODY))
            if business_value:
                elems.append(Paragraph(f"<b>Business Value:</b> {business_value}", STYLE_BODY))
            if criteria:
                elems.append(Paragraph("<b>Acceptance Criteria:</b>", STYLE_BODY))
                items = [ListItem(Paragraph(c, STYLE_BULLET)) for c in criteria]
                elems.append(ListFlowable(items, bulletType="bullet", start="- "))
            elems.append(Spacer(1, 6))

            story.append(KeepTogether(elems))

    # Version Plan
    if version_plan:
        story.append(Paragraph("Version Plan", STYLE_H2))
        for vp in version_plan:
            ver = vp.get("version", "Unknown")
            theme = vp.get("theme", "")
            release_goal = vp.get("release_goal", "")
            req_ids: List[str] = vp.get("requirement_ids") or []
            rationale = vp.get("rationale", "")

            elems = [
                Paragraph(f"v{ver} — {theme}", STYLE_H3),
                Paragraph(f"<b>Goal:</b> {release_goal}", STYLE_BODY),
                Paragraph(f"<b>Requirements:</b> {', '.join(req_ids) if req_ids else 'None'}", STYLE_BODY),
                Paragraph(f"<b>Rationale:</b> {rationale}", STYLE_BODY),
                Spacer(1, 4),
            ]
            story.append(KeepTogether(elems))

    # Footer
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#e2e8f0"),
        spaceBefore=16, spaceAfter=6))
    story.append(Paragraph(f"Generated by App Review Insights on {generated}", STYLE_FOOTER))

    return story


def export_to_pdf(prd: Dict[str, Any]) -> bytes:
    """Render a PRD dict to a PDF (bytes) via reportlab."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        title=prd.get("title", "PRD"),
        author="App Review Insights",
    )
    story = _build_pdf_story(prd)
    doc.build(story)
    buf.seek(0)
    return buf.read()


# ──────────────────────────────────────────────
# API Endpoint
# ──────────────────────────────────────────────

_CONTENT_TYPES = {
    "md": "text/markdown; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}

_EXTENSIONS = {"md": ".md", "docx": ".docx", "pdf": ".pdf"}


@router.post("/prd")
def export_prd(
    prd: Dict[str, Any],
    format: str = Query("pdf", description="Export format: md, docx, or pdf"),
) -> Response:
    """Export a PRD document in the requested format.

    Accepts the PRD JSON object in the request body.
    Supported formats: md (Markdown), docx (Word), pdf (PDF).

    Returns:
        FileResponse with appropriate Content-Type and Content-Disposition.
    """
    format = format.lower().strip()
    if format not in ("md", "docx", "pdf"):
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported format '{format}'. Choose md, docx, or pdf.",
        )

    app_name = _sanitize_filename(prd.get("app_name") or "export")
    date_str = _format_date(prd.get("generated_at", datetime.now()))[:10].replace("-", "")
    filename = f"PRD_{app_name}_{date_str}{_EXTENSIONS[format]}"

    try:
        if format == "md":
            content = export_to_markdown(prd)
            return Response(
                content=content.encode("utf-8"),
                media_type=_CONTENT_TYPES["md"],
                headers={"Content-Disposition": f'attachment; filename="{filename}"'},
            )
        elif format == "docx":
            content = export_to_docx(prd)
            return Response(
                content=content,
                media_type=_CONTENT_TYPES["docx"],
                headers={"Content-Disposition": f'attachment; filename="{filename}"'},
            )
        else:  # pdf
            content = export_to_pdf(prd)
            return Response(
                content=content,
                media_type=_CONTENT_TYPES["pdf"],
                headers={"Content-Disposition": f'attachment; filename="{filename}"'},
            )
    except Exception as e:
        logger.error("Export failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Export failed: {e}")
